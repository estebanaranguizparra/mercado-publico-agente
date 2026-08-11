#!/usr/bin/env python
"""Extrae el texto de un PDF o DOCX descargado, para poder leerlo y buscar en él.

Existe para las bases de licitación, que solo se publican como documentos
adjuntos. Usa librerías de Python puro (`pypdf`, `python-docx`) en vez de
`pdftotext`/`pandoc`, para no depender de binarios del sistema.

    .venv/bin/python scripts/extraer_texto.py bases.pdf
    .venv/bin/python scripts/extraer_texto.py bases.pdf --salida bases.txt

Un PDF escaneado no tiene capa de texto y devolverá poco o nada. En ese caso el
script lo dice explícitamente en vez de entregar un resultado vacío que se
podría confundir con "no dice nada al respecto": la diferencia importa, porque
un documento ilegible es un riesgo, no una ausencia de requisitos.
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

# Bajo este promedio de caracteres por página, el PDF casi con certeza es una
# imagen escaneada y no un documento con texto seleccionable.
_MINIMO_POR_PAGINA = 60


def _de_pdf(ruta: Path) -> tuple[str, int]:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise SystemExit(
            "Falta pypdf. Instálalo con: .venv/bin/pip install pypdf python-docx"
        )

    lector = PdfReader(str(ruta))
    partes = []
    for numero, pagina in enumerate(lector.pages, start=1):
        try:
            texto = pagina.extract_text() or ""
        except Exception as exc:  # noqa: BLE001 — una página rota no anula el resto
            texto = ""
            print(f"aviso: falló la página {numero}: {exc}", file=sys.stderr)
        partes.append(f"\n--- página {numero} ---\n{texto}")
    return "\n".join(partes), len(lector.pages)


def _de_docx(ruta: Path) -> tuple[str, int]:
    try:
        import docx
    except ImportError:
        raise SystemExit(
            "Falta python-docx. Instálalo con: .venv/bin/pip install pypdf python-docx"
        )

    documento = docx.Document(str(ruta))
    partes = [p.text for p in documento.paragraphs]

    # Las bases suelen poner las garantías y los criterios de evaluación en
    # tablas; sin esto se pierde justo lo que interesa.
    for tabla in documento.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells]
            if any(celdas):
                partes.append(" | ".join(celdas))

    return "\n".join(partes), len(documento.paragraphs)


def extraer(ruta: Path) -> tuple[str, str]:
    """Devuelve (texto, advertencia). La advertencia es '' si todo salió bien."""
    sufijo = ruta.suffix.lower()
    if sufijo == ".pdf":
        texto, paginas = _de_pdf(ruta)
        util = len(texto.replace("\n", "").strip())
        if paginas and util / paginas < _MINIMO_POR_PAGINA:
            return texto, (
                f"DOCUMENTO SIN TEXTO EXTRAÍBLE: {util} caracteres en {paginas} página(s). "
                "Casi con certeza es un escaneo. Necesita OCR (no disponible en este "
                "entorno): repórtalo como riesgo documental, no como ausencia de requisitos."
            )
        return texto, ""

    if sufijo in {".docx", ".doc"}:
        if sufijo == ".doc":
            return "", (
                "Formato .doc antiguo no soportado por python-docx. "
                "Ábrelo a mano o pide la versión .docx/.pdf."
            )
        texto, _ = _de_docx(ruta)
        return texto, ""

    if sufijo in {".txt", ".md"}:
        return ruta.read_text(encoding="utf-8", errors="replace"), ""

    return "", f"Formato no soportado: {sufijo}. Soporta .pdf, .docx, .txt"


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Extrae texto de un PDF o DOCX.")
    parser.add_argument("archivo", help="Ruta al documento.")
    parser.add_argument("--salida", help="Dónde escribir el texto (por defecto, stdout).")
    args = parser.parse_args(argv)

    ruta = Path(args.archivo)
    if not ruta.exists():
        print(f"No existe: {ruta}", file=sys.stderr)
        return 1

    texto, advertencia = extraer(ruta)

    if advertencia:
        print(f"⚠ {advertencia}", file=sys.stderr)

    if args.salida:
        destino = Path(args.salida)
        destino.parent.mkdir(parents=True, exist_ok=True)
        destino.write_text(texto, encoding="utf-8")
        print(f"{len(texto):,} caracteres → {destino}")
    else:
        print(texto)

    return 2 if advertencia else 0


if __name__ == "__main__":
    sys.exit(main())
